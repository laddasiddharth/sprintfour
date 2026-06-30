import io
import re
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse

router = APIRouter()

SUPPORTED_PLAIN = {"txt", "md", "markdown", "csv", "tsv", "log", "json", "xml"}
SUPPORTED_HTML = {"html", "htm"}


def strip_html_tags(html: str) -> str:
    """Strip HTML tags and decode common entities."""
    # Remove style and script blocks
    html = re.sub(r"<style[^>]*>[\s\S]*?</style>", " ", html, flags=re.IGNORECASE)
    html = re.sub(r"<script[^>]*>[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
    # Block-level newlines
    html = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    html = re.sub(r"</p>", "\n", html, flags=re.IGNORECASE)
    html = re.sub(r"</div>", "\n", html, flags=re.IGNORECASE)
    html = re.sub(r"</h[1-6]>", "\n", html, flags=re.IGNORECASE)
    html = re.sub(r"</li>", "\n", html, flags=re.IGNORECASE)
    # Remove remaining tags
    html = re.sub(r"<[^>]+>", " ", html)
    # Decode entities
    html = html.replace("&nbsp;", " ")
    html = html.replace("&amp;", "&")
    html = html.replace("&lt;", "<")
    html = html.replace("&gt;", ">")
    html = html.replace("&quot;", '"')
    html = html.replace("&#39;", "'")
    # Collapse whitespace
    html = re.sub(r"[ \t]{2,}", " ", html)
    html = re.sub(r"\n{3,}", "\n\n", html)
    return html.strip()


def strip_rtf(rtf: str) -> str:
    """Basic RTF stripping — removes control words and groups."""
    text = re.sub(r"\{\\[^{}]*\}", "", rtf)   # remove groups like {\fonttbl...}
    text = re.sub(r"\\[a-z]+\d*\s?", "", text) # remove control words like \rtf1 \par
    text = re.sub(r"[{}]", "", text)            # remove remaining braces
    return text.strip()


@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    filename = (file.filename or "").lower()
    ext = filename.rsplit(".", 1)[-1] if "." in filename else ""
    content = await file.read()

    text = ""

    # ── Plain text formats ──────────────────────────────────────────────────
    if ext in SUPPORTED_PLAIN:
        text = content.decode("utf-8", errors="replace")

    # ── HTML ────────────────────────────────────────────────────────────────
    elif ext in SUPPORTED_HTML:
        text = strip_html_tags(content.decode("utf-8", errors="replace"))

    # ── PDF ─────────────────────────────────────────────────────────────────
    elif ext == "pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                text = "\n\n".join(pages)
            if not text.strip():
                raise HTTPException(
                    status_code=422,
                    detail="This PDF appears to be a scanned image. Text could not be extracted. Please use a text-based PDF.",
                )
        except HTTPException:
            raise
        except Exception as err:
            raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {err}")

    # ── Word Documents (.docx, .doc) ────────────────────────────────────────
    elif ext in {"docx", "doc"}:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs]
            # Also include table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            paragraphs.append(para.text)
            text = "\n".join(paragraphs)
        except Exception as err:
            raise HTTPException(status_code=500, detail=f"Failed to extract text from Word document: {err}")

    # ── RTF ──────────────────────────────────────────────────────────────────
    elif ext == "rtf":
        text = strip_rtf(content.decode("utf-8", errors="replace"))

    # ── Unsupported ──────────────────────────────────────────────────────────
    else:
        raise HTTPException(
            status_code=415,
            detail=f'Unsupported file type ".{ext}". Supported formats: PDF, DOCX, DOC, TXT, MD, CSV, HTML, RTF.',
        )

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="The file appears to be empty or contains no readable text.",
        )

    return {"text": text}
