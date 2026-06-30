"""
POST /api/redact-doc

Accepts:
  - file: UploadFile  — the original document (DOCX, TXT, MD, HTML, etc.)
  - redactions: str   — JSON array of { text: string, type: string }

Returns the same document format with PII text replaced by black bars
(█ full-block characters). No [TAG] placeholders — purely visual black bars.
"""

import io
import json
import re
from typing import Optional
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse

router = APIRouter()

BLOCK_CHAR = "█"  # U+2588 FULL BLOCK — renders as a solid black bar

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def build_replacement_map(redactions: list[dict]) -> dict[str, str]:
    """
    Build {original_text: '█████'} mapping.
    Replacement is the same character count as the original text (so the
    visual width stays approximately the same).
    Sorted longest-first to avoid partial-match issues.
    """
    mapping: dict[str, str] = {}
    for r in redactions:
        text = r.get("text", "").strip()
        if text:
            # Replace with same number of █ characters for same visual weight
            mapping[text] = BLOCK_CHAR * len(text)
    return dict(sorted(mapping.items(), key=lambda kv: len(kv[0]), reverse=True))


def apply_replacements(text: str, mapping: dict[str, str]) -> str:
    for original, replacement in mapping.items():
        text = text.replace(original, replacement)
    return text


# ─────────────────────────────────────────────────────────────────────────────
# DOCX: replace text with █ blocks, preserve paragraph/run formatting
# ─────────────────────────────────────────────────────────────────────────────

def _redact_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """
    Redact PII in a single DOCX paragraph.
    Merges all runs into the first run (preserving paragraph-level style),
    replaces PII with █ blocks.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(original in full_text for original in mapping):
        return

    new_text = apply_replacements(full_text, mapping)
    if new_text == full_text:
        return

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    # Write redacted text into first run, blank out the rest
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def redact_docx(content: bytes, mapping: dict[str, str]) -> bytes:
    from docx import Document

    doc = Document(io.BytesIO(content))

    for para in doc.paragraphs:
        _redact_paragraph(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _redact_paragraph(para, mapping)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _redact_paragraph(para, mapping)
        for para in section.footer.paragraphs:
            _redact_paragraph(para, mapping)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# HTML: replace text nodes with █ blocks
# ─────────────────────────────────────────────────────────────────────────────

def redact_html(content: bytes, mapping: dict[str, str]) -> bytes:
    from bs4 import BeautifulSoup, NavigableString

    soup = BeautifulSoup(content.decode("utf-8", errors="replace"), "html.parser")

    def _replace_in_node(node):
        if isinstance(node, NavigableString):
            new_text = apply_replacements(str(node), mapping)
            if new_text != str(node):
                node.replace_with(new_text)
        else:
            for child in list(node.children):
                _replace_in_node(child)

    _replace_in_node(soup)
    return str(soup).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# Plain text / RTF: simple string replacement
# ─────────────────────────────────────────────────────────────────────────────

def redact_plain(content: bytes, mapping: dict[str, str]) -> bytes:
    text = content.decode("utf-8", errors="replace")
    return apply_replacements(text, mapping).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# Route
# ─────────────────────────────────────────────────────────────────────────────

PLAIN_EXTS = {"txt", "md", "markdown", "csv", "tsv", "log", "json", "xml", "rtf"}
HTML_EXTS = {"html", "htm"}

MIME_MAP = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "md": "text/plain",
    "markdown": "text/plain",
    "csv": "text/csv",
    "tsv": "text/tab-separated-values",
    "html": "text/html",
    "htm": "text/html",
    "json": "application/json",
    "xml": "application/xml",
    "log": "text/plain",
    "rtf": "application/rtf",
}


@router.post("/redact-doc")
async def redact_doc(
    file: UploadFile = File(...),
    redactions: str = Form(...),
):
    filename = (file.filename or "document").lower()
    ext = filename.rsplit(".", 1)[-1] if "." in filename else "txt"
    content = await file.read()

    try:
        redaction_list: list[dict] = json.loads(redactions)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON in 'redactions' field.")

    if not redaction_list:
        return StreamingResponse(
            io.BytesIO(content),
            media_type=MIME_MAP.get(ext, "application/octet-stream"),
            headers={"Content-Disposition": f'attachment; filename="redacted-{file.filename}"'},
        )

    mapping = build_replacement_map(redaction_list)

    try:
        if ext in {"docx", "doc"}:
            redacted_bytes = redact_docx(content, mapping)
        elif ext in HTML_EXTS:
            redacted_bytes = redact_html(content, mapping)
        else:
            # TXT, MD, CSV, RTF, etc.
            redacted_bytes = redact_plain(content, mapping)

    except Exception as err:
        raise HTTPException(status_code=500, detail=f"Failed to redact document: {err}")

    mime = MIME_MAP.get(ext, "application/octet-stream")
    safe_name = file.filename or f"document.{ext}"

    return StreamingResponse(
        io.BytesIO(redacted_bytes),
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="redacted-{safe_name}"'},
    )
